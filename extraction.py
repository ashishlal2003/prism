import PyPDF2
import os
from tqdm import tqdm
import docx

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file"""

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found at: {pdf_path}")

    print(f"Extracting text from: {os.path.basename(pdf_path)}")

    text = ""

    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            print(f"PDF has {num_pages} pages")

            for page_num in tqdm(range(num_pages), desc="Processing pages"):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"

            print("Text extraction complete!")

    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file"""

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found at: {docx_path}")

    print(f"Extracting text from: {os.path.basename(docx_path)}")

    text = ""

    try:
        doc = docx.Document(docx_path)

        for para in tqdm(doc.paragraphs, desc="Processing paragraphs"):
            if para.text:
                text += para.text + "\n\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
            text += "\n\n"

        print("Text extraction complete!")

    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None

    return text